import streamlit as st
import docx
from io import BytesIO
from sentence_transformers import SentenceTransformer
from difflib import SequenceMatcher

# Load pre-trained sentence transformer model
model = SentenceTransformer('paraphrase-MiniLM-L6-v2')

# Function to extract text from Word documents
def extract_text_from_docx(doc):
    return ' '.join([paragraph.text for paragraph in doc.paragraphs])

# Function to chunk text into paragraphs
def chunk_text(text):
    return text.split("\n\n")

# Function to calculate similarity
def calculate_similarity(paragraph1, paragraph2):
    embedding1 = model.encode(paragraph1)
    embedding2 = model.encode(paragraph2)
    similarity = SequenceMatcher(None, paragraph1, paragraph2).ratio()
    return similarity

# Streamlit UI
st.title("Legal Document Comparison and Adaptation")

# Upload source and target Word documents
source_file = st.file_uploader("Upload Source Document", type=["docx"])
target_file = st.file_uploader("Upload Target Document", type=["docx"])

if source_file and target_file:
    source_doc = docx.Document(source_file)
    target_doc = docx.Document(target_file)

    source_text = extract_text_from_docx(source_doc)
    target_text = extract_text_from_docx(target_doc)

    source_paragraphs = chunk_text(source_text)
    target_paragraphs = chunk_text(target_text)

    if len(source_paragraphs) != len(target_paragraphs):
        st.warning("The documents have different numbers of paragraphs. Comparison may not be accurate.")
    
    # Button to compare
    if st.button("Compare"):
        st.write("### Comparison Results")
        for i, (source_paragraph, target_paragraph) in enumerate(zip(source_paragraphs, target_paragraphs)):
            similarity = calculate_similarity(source_paragraph, target_paragraph)
            st.write(f"**Paragraph {i+1}: Similarity: {similarity:.2f}**")
            st.write(f"Source: {source_paragraph}")
            st.write(f"Target: {target_paragraph}")
            if similarity < 0.8:
                st.write(f":red[Changes detected]")
            st.write("---")
    
    # Button to adapt
    if st.button("Adapt"):
        adapted_doc = docx.Document()
        for source_paragraph, target_paragraph in zip(source_paragraphs, target_paragraphs):
            similarity = calculate_similarity(source_paragraph, target_paragraph)
            if similarity < 0.8:
                adapted_doc.add_paragraph(source_paragraph)
            else:
                adapted_doc.add_paragraph(target_paragraph)

        # Save the adapted document
        adapted_io = BytesIO()
        adapted_doc.save(adapted_io)
        adapted_io.seek(0)
        st.download_button("Download Adapted Document", data=adapted_io, file_name="adapted.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

